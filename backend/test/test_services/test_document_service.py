from fastapi import UploadFile
from starlette.datastructures import Headers
from app.services.document_service import DocumentService
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
import pytest
import os

from app.exceptions.custom_exceptions import (
    InvalidFileExtensionError,
    InvalidMimeTypeError,
    ConversionError
)

service = DocumentService()

@pytest.fixture()
def txt_file():
    content = """"
    Arquivo TXT de teste

    linha 1
    linha 2
    linha 3
    """

    return BytesIO(content.encode("utf-8"))

@pytest.fixture()
def docx_file():
    document = Document()

    document.add_heading("Arquivo DOCX de teste", level=1)
    document.add_paragraph("Docuemtno de teste para conversão")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer

@pytest.fixture()
def pdf_file():
    buffer = BytesIO()

    pdf = canvas.Canvas(buffer)
    pdf.drawString(
        100,
        750,
        "PDF de teste para conversão"
    )

    pdf.save()
    buffer.seek(0)

    return buffer

class TestDocumentService:
    #TESTS FOR SERVICE convert_txt_to_pdf
    @pytest.mark.asyncio
    async def test_txt_to_pdf_success(self, txt_file):
        txt = txt_file

        upload_file = UploadFile(
            filename="test.txt",
            file=txt,
            headers=Headers({
                "content-type": "text/plain"
            })
        )

        response = await service.convert_txt_to_pdf(upload_file)

        assert response.endswith(".pdf")
        assert os.path.exists(response)

    @pytest.mark.asyncio
    async def test_txt_to_pdf_conversion_error(self, txt_file, mocker):
        txt = txt_file

        upload_file = UploadFile(
            filename="test.txt",
            file=txt,
            headers=Headers({
                "content-type": "text/plain"
            })
        )

        mocker.patch(
            "app.services.document_service.canvas.Canvas.save",
            side_effect=Exception("Erro simulado na conversão")
        )

        with pytest.raises(ConversionError) as exc:
            await service.convert_txt_to_pdf(upload_file)

        assert str(exc.value).startswith("Erro ao converter TXT para PDF: ")

    @pytest.mark.asyncio
    async def test_txt_to_pdf_invalid_extension(self, txt_file):
        txt = txt_file

        upload_file = UploadFile(
            filename="test.docx",
            file=txt,
            headers=Headers({
                "content-type": "text/plain"
            })
        )

        with pytest.raises(InvalidFileExtensionError) as exc:
            await service.convert_txt_to_pdf(upload_file)

        assert str(exc.value) == "Extensão inválida. Extensões permitidas: .txt"

    @pytest.mark.asyncio
    async def test_txt_to_pdf_invalid_mime_type(self, txt_file):
        txt = txt_file

        upload_file = UploadFile(
            filename="test.txt",
            file=txt,
            headers=Headers({
                "content-type": "application/pdf"
            })
        )

        with pytest.raises(InvalidMimeTypeError) as exc:
            await service.convert_txt_to_pdf(upload_file)
        
        assert str(exc.value) == (
            "Tipo MIME inválido. Tipos MIME permitidos: text/plain"
        )

    #TESTS FOR SERVICE convert_txt_to_docx
    @pytest.mark.asyncio
    async def test_txt_to_docx_success(self, txt_file):
        txt = txt_file

        upload_file = UploadFile(
            filename="test.txt",
            file=txt,
            headers=Headers({
                "content-type": "text/plain"
            })
        )

        response = await service.convert_txt_to_docx(upload_file)

        assert response.endswith(".docx")
        assert os.path.exists(response)

    @pytest.mark.asyncio
    async def test_txt_to_docx_conversion_error(self, txt_file, mocker):
        txt = txt_file

        upload_file = UploadFile(
            filename="test.txt",
            file=txt,
            headers=Headers({
                "content-type": "text/plain"
            })
        )

        mocker.patch(
            "app.services.document_service.Document",
            side_effect=Exception("Erro simulado na conversão")
        )

        with pytest.raises(ConversionError) as exc:
            await service.convert_txt_to_docx(upload_file)

        assert str(exc.value).startswith("Erro ao converter TXT para DOCX:")

    @pytest.mark.asyncio
    async def test_txt_to_docx_invalid_extension(self, txt_file):
        txt = txt_file

        upload_file = UploadFile(
            filename="test.pdf",
            file=txt,
            headers=Headers({
                "content-type": "text/plain"
            })
        )

        with pytest.raises(InvalidFileExtensionError) as exc:
            await service.convert_txt_to_docx(upload_file)

        assert str(exc.value) == "Extensão inválida. Extensões permitidas: .txt"

    @pytest.mark.asyncio
    async def test_txt_to_docx_invalid_mime_type(self, txt_file):
        txt = txt_file

        upload_file = UploadFile(
            filename="test.txt",
            file=txt,
            headers=Headers({
                "content-type": "application/pdf"
            })
        )

        with pytest.raises(InvalidMimeTypeError) as exc:
            await service.convert_txt_to_docx(upload_file)
        
        assert str(exc.value) == (
            "Tipo MIME inválido. Tipos MIME permitidos: text/plain"
        )

    #TESTS FOR SERVICE convert_docx_to_pdf
    @pytest.mark.asyncio
    async def test_docx_to_pdf_success(self, docx_file):
        docx = docx_file

        upload_file = UploadFile(
            filename="test.docx",
            file=docx,
            headers=Headers({
                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            })
        )

        response = await service.convert_docx_to_pdf(upload_file)

        assert response["output_path"].endswith(".pdf")
        assert os.path.exists(response["output_path"])

    @pytest.mark.asyncio
    async def test_docx_to_pdf_conversion_error(self, docx_file, mocker):
        docx = docx_file

        upload_file = UploadFile(
            filename="test.docx",
            file=docx,
            headers=Headers({
                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            })
        )

        mocker.patch(
            "app.services.document_service.subprocess.run",
            side_effect=Exception("Erro simulado na conversão")
        )

        with pytest.raises(ConversionError) as exc:
            await service.convert_docx_to_pdf(upload_file)

        assert str(exc.value).startswith("Erro ao converter DOCX para PDF: ")

    @pytest.mark.asyncio
    async def test_docx_to_pdf_invalid_extension(self, docx_file):
        docx = docx_file

        upload_file = UploadFile(
            filename="test.txt",
            file=docx,
            headers=Headers({
                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            })
        )

        with pytest.raises(InvalidFileExtensionError) as exc:
            await service.convert_docx_to_pdf(upload_file)

        assert str(exc.value) == (
            "Extensão inválida. Extensões permitidas: .docx"
        )

    @pytest.mark.asyncio
    async def test_docx_to_pdf_invalid_mime_type(self, docx_file):
        docx = docx_file

        upload_file = UploadFile(
            filename="test.docx",
            file=docx,
            headers=Headers({
                "content-type": "text/plain"
            })
        )

        with pytest.raises(InvalidMimeTypeError) as exc:
            await service.convert_docx_to_pdf(upload_file)
        
        assert str(exc.value) == (
            "Tipo MIME inválido. Tipos MIME permitidos: application/vnd.openxmlformats-officedocument.wordprocessingml.document, application/zip"
        )

    #TESTS FOR SERVICE convert_pdf_to_docx
    @pytest.mark.asyncio
    async def test_pdf_to_docx_success(self, pdf_file):
        pdf = pdf_file

        upload_file = UploadFile(
            filename="test.pdf",
            file=pdf,
            headers=Headers({
                "content-type": "application/pdf"
            })
        )

        response = await service.convert_pdf_to_docx(upload_file)

        assert response["output_path"].endswith(".docx")
        assert os.path.exists(response["output_path"])

    @pytest.mark.asyncio
    async def test_pdf_to_docx_conversion_error(self, pdf_file, mocker):
        pdf = pdf_file

        upload_file = UploadFile(
            filename="test.pdf",
            file=pdf,
            headers=Headers({
                "content-type": "application/pdf"
            })
        )

        mocker.patch(
            "app.services.document_service.pdfplumber.open",
            side_effect=Exception("Erro simulado na conversão")
        )

        with pytest.raises(ConversionError) as exc:
            await service.convert_pdf_to_docx(upload_file)

        assert str(exc.value).startswith("Erro ao converter PDF para DOCX: ")

    @pytest.mark.asyncio
    async def test_pdf_to_docx_invalid_extension(self, pdf_file):
        pdf = pdf_file

        upload_file = UploadFile(
            filename="test.txt",
            file=pdf,
            headers=Headers({
                "content-type": "application/pdf"
            })
        )

        with pytest.raises(InvalidFileExtensionError) as exc:
            await service.convert_pdf_to_docx(upload_file)

        assert str(exc.value) == (
            "Extensão inválida. Extensões permitidas: .pdf"
        )

    @pytest.mark.asyncio
    async def test_pdf_to_docx_invalid_mime_type(self, pdf_file):
        pdf = pdf_file

        upload_file = UploadFile(
            filename="test.pdf",
            file=pdf,
            headers=Headers({
                "content-type": "text/plain"
            })
        )

        with pytest.raises(InvalidMimeTypeError) as exc:
            await service.convert_pdf_to_docx(upload_file)
        
        assert str(exc.value) == (
            "Tipo MIME inválido. Tipos MIME permitidos: application/pdf"
        )