from app.exceptions.custom_exceptions import ConversionError
from app.core.utils import ensure_output_dir, get_output_path, get_input_path
from app.core.validators import validate_file
from reportlab.pdfgen import canvas
from docx import Document
import os
import pdfplumber
import subprocess

class DocumentService:
    async def convert_txt_to_pdf(self, file):
        validate_file(file, "txt")

        try:
            output_dir = ensure_output_dir()
            output_path = get_output_path(file, output_dir, "pdf")

            content = (await file.read()).decode("utf-8")

            pdf = canvas.Canvas(output_path)
            pdf.drawString(100, 750, content)
            pdf.save()

            return output_path

        except Exception as e:
            raise ConversionError(
                f"Erro ao converter TXT para PDF: {str(e)}"
            ) from e

    async def convert_txt_to_docx(self, file):
        validate_file(file, "txt")

        try:
            output_dir = ensure_output_dir()
            output_path = get_output_path(file, output_dir, "docx")

            content = (await file.read()).decode("utf-8")

            doc = Document()
            doc.add_paragraph(content)
            doc.save(output_path)

            return output_path

        except Exception as e:
            raise ConversionError(
                f"Erro ao converter TXT para DOCX: {str(e)}"
            ) from e

    async def convert_docx_to_pdf(self, file):
        validate_file(file, "docx")

        try:
            output_dir = ensure_output_dir()
            input_path = get_input_path(file, output_dir)
            output_path = get_output_path(file, output_dir, "pdf")

            with open(input_path, "wb") as f:
                f.write(await file.read())

            subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    input_path,
                    "--outdir",
                    output_dir,
                ],
                check=True,
            )

            return {
                "output_path": output_path,
                "input_path": input_path,
            }
        
        except subprocess.CalledProcessError as e:
            raise ConversionError(
                "Erro ao converter DOCX para PDF"
            ) from e
        
        except FileNotFoundError as e:
            raise ConversionError(
                "LibreOffice não encontrado no sistema"
            ) from e
        
        except Exception as e:
            raise ConversionError(
                f"Erro ao converter DOCX para PDF: {str(e)}"
            ) from e

    async def convert_pdf_to_docx(self, file):
        validate_file(file, "pdf")

        try:
            output_dir = ensure_output_dir()
            input_path = get_input_path(file, output_dir)
            output_path = get_output_path(file, output_dir, "docx")

            with open(input_path, "wb") as f:
                f.write(await file.read())

            doc = Document()

            with pdfplumber.open(input_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()

                    if text:
                        for line in text.splitlines():
                            doc.add_paragraph(line)
            
            doc.save(output_path)
            return {
                "output_path": output_path,
                "input_path": input_path
            } 
    
        except Exception as e:
            raise ConversionError(
                f"Erro ao converter PDF para DOCX: {str(e)}"
            ) from e
